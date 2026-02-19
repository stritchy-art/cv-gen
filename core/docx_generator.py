"""
Module de génération de fichiers DOCX
Crée un fichier Word formaté avec le même style visuel que CV_exemple.html
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from config.logging_config import setup_logger

# Removed unused imports - types used via dict, list builtin types


# Logger
logger = setup_logger(__name__, "docx_generator.log")


class CVDocxGenerator:
    """Générateur de CV au format DOCX avec style personnalisé"""

    # Couleurs définies selon CV_exemple.html
    COLOR_DARK_BLUE = RGBColor(29, 67, 91)  # #1D435B
    COLOR_GOLD = RGBColor(188, 148, 74)  # #BC944A
    COLOR_LIGHT_GOLD = RGBColor(171, 141, 83)  # #AB8D53
    COLOR_BLACK = RGBColor(0, 0, 0)
    COLOR_DARK_GRAY = RGBColor(51, 51, 51)
    COLOR_GRAY = RGBColor(68, 68, 68)
    COLOR_LIGHT_GRAY = RGBColor(192, 192, 192)

    # Traductions des labels du CV
    LABELS = {
        "fr": {
            "competences": "Compétences",
            "competences_op": "Compétences opérationnelles",
            "competences_tech": "Compétences techniques",
            "formations": "Formations",
            "experiences": "Expériences professionnelles",
            "contexte": "Contexte",
            "activites": "Activités",
            "env_tech": "Environnement technique",
        },
        "en": {
            "competences": "Skills",
            "competences_op": "Operational Skills",
            "competences_tech": "Technical Skills",
            "formations": "Education",
            "experiences": "Professional Experience",
            "contexte": "Context",
            "activites": "Activities",
            "env_tech": "Technical Environment",
        },
        "it": {
            "competences": "Competenze",
            "competences_op": "Competenze Operative",
            "competences_tech": "Competenze Tecniche",
            "formations": "Formazione",
            "experiences": "Esperienza Professionale",
            "contexte": "Contesto",
            "activites": "Attività",
            "env_tech": "Ambiente Tecnico",
        },
        "es": {
            "competences": "Competencias",
            "competences_op": "Competencias Operativas",
            "competences_tech": "Competencias Técnicas",
            "formations": "Formación",
            "experiences": "Experiencia Profesional",
            "contexte": "Contexto",
            "activites": "Actividades",
            "env_tech": "Entorno Técnico",
        },
    }

    def __init__(self, target_language="fr"):
        self.doc = Document()
        self.target_language = (
            target_language if target_language in self.LABELS else "fr"
        )
        self.labels = self.LABELS[self.target_language]
        self._setup_document()
        self._setup_numbering()

    def _setup_document(self):
        """Configure les marges et le style du document"""
        # Marges personnalisées
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Cm(5.57)  # 5.57 cm en haut
            section.bottom_margin = Cm(0.49)  # 0.49 cm en bas
            section.left_margin = Cm(1.76)  # 1.76 cm à gauche
            section.right_margin = Cm(1.76)  # 1.76 cm à droite

    def _add_page_footer(self):
        """Ajoute un pied de page avec numérotation, date et copyright"""
        from datetime import datetime

        from docx.oxml import parse_xml, register_element_cls
        from docx.oxml.ns import nsdecls

        section = self.doc.sections[0]
        footer = section.footer

        # Créer un tableau avec 2 colonnes pour le pied de page
        table = footer.add_table(rows=1, cols=2, width=Inches(6.88))
        table.autofit = False
        table.allow_autofit = False

        # Définir les largeurs des colonnes
        table.columns[0].width = Inches(3.44)  # 50% pour la gauche
        table.columns[1].width = Inches(3.44)  # 50% pour la droite

        # Supprimer les bordures du tableau
        for row in table.rows:
            for cell in row.cells:
                cell._element.get_or_add_tcPr().append(
                    parse_xml(
                        r'<w:tcBorders %s><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/></w:tcBorders>'
                        % nsdecls("w")
                    )
                )

        # Cellule gauche : "Page n/n"
        left_cell = table.rows[0].cells[0]
        left_para = left_cell.paragraphs[0]
        left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Ajouter "Page "
        run = left_para.add_run("Page ")
        run.font.size = Pt(9)
        run.font.color.rgb = self.COLOR_GRAY

        # Ajouter le numéro de page (field code)
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        run._r.append(fldChar1)

        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = "PAGE"
        run._r.append(instrText)

        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar2)

        # Ajouter "/"
        run = left_para.add_run("/")
        run.font.size = Pt(9)
        run.font.color.rgb = self.COLOR_GRAY

        # Ajouter le nombre total de pages (field code)
        fldChar3 = OxmlElement("w:fldChar")
        fldChar3.set(qn("w:fldCharType"), "begin")
        run._r.append(fldChar3)

        instrText2 = OxmlElement("w:instrText")
        instrText2.set(qn("xml:space"), "preserve")
        instrText2.text = "NUMPAGES"
        run._r.append(instrText2)

        fldChar4 = OxmlElement("w:fldChar")
        fldChar4.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar4)

        # Cellule droite : Date et copyright
        right_cell = table.rows[0].cells[1]
        right_para = right_cell.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Date du jour
        date_str = datetime.now().strftime("%d/%m/%Y")
        run = right_para.add_run(f"{date_str} - ")
        run.font.size = Pt(9)
        run.font.color.rgb = self.COLOR_GRAY

        # Copyright
        run = right_para.add_run(
            "© ALLTECH - Tous droits réservés / CONFIDENTIEL ALLTECH"
        )
        run.font.size = Pt(9)
        run.font.color.rgb = self.COLOR_GRAY
        run.bold = True

    def _setup_numbering(self):
        """Configure le style de numérotation pour les puces carrées bleues"""
        # Créer la partie numbering si elle n'existe pas
        try:
            numbering_part = self.doc.part.numbering_part
        except:
            from docx.oxml.numbering import CT_Numbering

            numbering_part = self.doc.part.add_numbering_part()
            numbering_part.element = CT_Numbering()

        # Créer une définition de numérotation avec puce carrée
        numbering_xml = """
        <w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:abstractNum w:abstractNumId="0">
                <w:multiLevelType w:val="hybridMultilevel"/>
                <w:lvl w:ilvl="0">
                    <w:start w:val="1"/>
                    <w:numFmt w:val="bullet"/>
                    <w:lvlText w:val="■"/>
                    <w:lvlJc w:val="left"/>
                    <w:pPr>
                        <w:ind w:left="720" w:hanging="360"/>
                    </w:pPr>
                    <w:rPr>
                        <w:color w:val="1D435B"/>
                        <w:sz w:val="12"/>
                    </w:rPr>
                </w:lvl>
            </w:abstractNum>
            <w:num w:numId="1">
                <w:abstractNumId w:val="0"/>
            </w:num>
        </w:numbering>
        """

        from lxml import etree

        numbering_element = etree.fromstring(numbering_xml)

        # Ajouter les éléments au numbering part
        for child in numbering_element:
            numbering_part.element.append(child)

    def _add_page_header(self, cv_data, output_path):
        """Ajoute une en-tête de page avec logo, 'Fiche de compétences' et nom"""
        header = cv_data.get("header", {})

        # Créer l'en-tête de page
        section = self.doc.sections[0]
        header_section = section.header

        # Créer un tableau avec 3 colonnes pour un meilleur équilibre visuel
        from docx.shared import Inches
        from docx.table import Table

        table = header_section.add_table(rows=1, cols=3, width=Inches(6.88))
        table.autofit = False
        table.allow_autofit = False

        # Définir les largeurs des colonnes pour centrer visuellement le texte et assurer une seule ligne
        table.columns[0].width = Inches(0.7)  # Logo (réduit)
        table.columns[1].width = Inches(
            6.0
        )  # Fiche de compétences + Nom (centré, très large)
        table.columns[2].width = Inches(0.18)  # Colonne vide réduite pour équilibrer

        cells = table.rows[0].cells

        # Colonne 1 : Logo
        # Chercher le logo dans le dossier assets
        script_dir = Path(__file__).parent
        logo_path = script_dir.parent / "assets" / "logo_alltech.png"
        if not logo_path.exists():
            # Fallback : dossier du script ou parent du output
            logo_path = script_dir / "logo_alltech.png"
            if not logo_path.exists():
                logo_path = Path(output_path).parent / "logo_alltech.png"

        if logo_path.exists():
            logo_para = cells[0].paragraphs[0]
            run = logo_para.add_run()
            run.add_picture(str(logo_path), width=Inches(0.8))  # Logo plus petit

        # Colonne 2 : "Fiche de compétences" + Nom centré sur une seule ligne
        fiche_para = cells[1].paragraphs[0]
        fiche_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Empêcher le retour à la ligne
        from docx.oxml.shared import OxmlElement

        pPr = fiche_para._element.get_or_add_pPr()
        noWrap = OxmlElement("w:noWrap")
        pPr.append(noWrap)

        # "Fiche de compétences" en bleu
        fiche_run = fiche_para.add_run("Fiche de compétences ")
        fiche_run.font.size = Pt(18)
        fiche_run.font.color.rgb = self.COLOR_DARK_BLUE
        fiche_run.font.bold = True
        fiche_run.font.name = "Arial MT"

        # Nom de la personne en marron juste à côté
        name_run = fiche_para.add_run(header.get("name", "").upper())
        name_run.font.size = Pt(18)
        name_run.font.color.rgb = self.COLOR_GOLD  # Marron/Doré
        name_run.font.bold = True
        name_run.font.name = "Arial MT"

        # Forcer la cellule à ne pas couper le texte
        tc = cells[1]._element
        tcPr = tc.get_or_add_tcPr()
        noWrap_cell = OxmlElement("w:noWrap")
        tcPr.append(noWrap_cell)

        # Colonne 3 : Vide (pour équilibrer visuellement avec le logo)

        # Supprimer les bordures du tableau sauf le bas (ligne de séparation)
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        tbl = table._element
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)

        tblBorders = OxmlElement("w:tblBorders")
        # Supprimer toutes les bordures sauf celle du bas
        for border_name in ["top", "left", "right", "insideH", "insideV"]:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "none")
            tblBorders.append(border)

        # Ajouter une bordure inférieure noire plus fine
        bottom_border = OxmlElement("w:bottom")
        bottom_border.set(qn("w:val"), "single")
        bottom_border.set(qn("w:sz"), "6")  # Épaisseur réduite
        bottom_border.set(qn("w:space"), "0")
        bottom_border.set(qn("w:color"), "000000")  # Noir
        tblBorders.append(bottom_border)

        tblPr.append(tblBorders)

    def _add_header(self, cv_data):
        """Ajoute l'en-tête du CV"""
        header = cv_data.get("header", {})

        # Titre (20pt, bleu foncé, centré, majuscules)
        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(header.get("title", "").upper())
        title_run.font.size = Pt(20)
        title_run.font.color.rgb = self.COLOR_DARK_BLUE
        title_run.font.bold = True
        title_run.font.name = "Calibri"

        # Expérience (20pt, doré, centré)
        exp_para = self.doc.add_paragraph()
        exp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        exp_run = exp_para.add_run(header.get("experience", ""))
        exp_run.font.size = Pt(20)
        exp_run.font.color.rgb = self.COLOR_LIGHT_GOLD
        exp_run.font.bold = True
        exp_run.font.name = "Calibri"

        # Espacement après l'en-tête
        exp_para.paragraph_format.space_after = Pt(20)

    def _add_section_title(self, title):
        """Ajoute un titre de section (fond gris clair, bleu foncé, italique)"""
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(30)
        para.paragraph_format.space_after = Pt(15)

        run = para.add_run(title.upper())
        run.font.size = Pt(12)
        run.font.color.rgb = self.COLOR_DARK_BLUE
        run.font.bold = True
        run.font.italic = True
        run.font.name = "Calibri"

        # Ajout du fond gris clair
        shading_elm = OxmlElement("w:shd")
        shading_elm.set(qn("w:fill"), "D3D3D3")  # lightgray
        para._element.get_or_add_pPr().append(shading_elm)

    def _add_subsection_title(self, title):
        """Ajoute un sous-titre de section (12pt, gras, souligné)"""
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(15)
        para.paragraph_format.space_after = Pt(8)

        run = para.add_run(title)
        run.font.size = Pt(12)
        run.font.color.rgb = self.COLOR_DARK_GRAY
        run.font.bold = True
        run.font.underline = True
        run.font.name = "Calibri"

    def _add_bullet_list(self, items):
        """Ajoute une liste à puces carrées bleues"""
        for item in items:
            para = self.doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(1.0)
            para.paragraph_format.space_after = Pt(4)

            # Configuration pour puce carrée (■)
            pPr = para._element.get_or_add_pPr()
            numPr = OxmlElement("w:numPr")

            # Utiliser le format de puce carré (lvl 0, numId 1)
            ilvl = OxmlElement("w:ilvl")
            ilvl.set(qn("w:val"), "0")
            numId = OxmlElement("w:numId")
            numId.set(qn("w:val"), "1")

            numPr.append(ilvl)
            numPr.append(numId)
            pPr.append(numPr)

            # Texte de la liste
            run = para.add_run(item)
            run.font.size = Pt(12)
            run.font.color.rgb = self.COLOR_DARK_GRAY
            run.font.name = "Calibri"

    def _add_competences(self, cv_data):
        """Ajoute la section Compétences"""
        competences = cv_data.get("competences", {})

        self._add_section_title(self.labels["competences"])

        # Compétences opérationnelles
        if competences.get("operationnelles"):
            self._add_subsection_title(self.labels["competences_op"])
            self._add_bullet_list(competences["operationnelles"])

        # Compétences techniques
        if competences.get("techniques"):
            self._add_subsection_title(self.labels["competences_tech"])
            tech_items = []
            for tech in competences["techniques"]:
                if isinstance(tech, dict):
                    category = tech.get("category", "")
                    items = tech.get("items", [])
                    # Formater : "Catégorie : item1, item2, item3"
                    if isinstance(items, list):
                        items_str = ", ".join(items)
                    else:
                        items_str = str(items)
                    tech_items.append(
                        f"{category} : {items_str}" if category else items_str
                    )
                else:
                    tech_items.append(str(tech))
            self._add_bullet_list(tech_items)

    def _add_formations(self, cv_data):
        """Ajoute la section Formations"""
        formations = cv_data.get("formations", [])

        if not formations:
            return

        self._add_section_title(self.labels["formations"])

        formation_items = []
        for formation in formations:
            if isinstance(formation, dict):
                year = formation.get("year", "")
                description = formation.get("description", "")
                formation_items.append(
                    f"{year} : {description}" if year else description
                )
            else:
                formation_items.append(str(formation))

        self._add_bullet_list(formation_items)

    def _add_experience_block(self, experience):
        """Ajoute un bloc d'expérience professionnelle"""
        # En-tête de l'expérience (entreprise et période sur une seule ligne)
        header_para = self.doc.add_paragraph()
        header_para.paragraph_format.space_before = Pt(10)
        header_para.paragraph_format.space_after = Pt(5)

        # Éviter que l'expérience commence en fin de page (page-break-inside: avoid)
        header_para.paragraph_format.keep_with_next = True
        header_para.paragraph_format.keep_together = True

        # Ajouter un taquet de tabulation à droite pour aligner la période
        # Largeur A4 (21cm) - marges gauche (1.76cm) - marges droite (1.76cm) = 17.48cm ≈ 6.88 pouces
        from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER

        tab_stops = header_para.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(17.48), WD_TAB_ALIGNMENT.RIGHT)

        # Entreprise (gauche, doré, gras, majuscules)
        # Nettoyer le nom de l'entreprise : enlever "(NON RENSEIGNÉ)" ou "()" vides
        import re

        company_name = experience.get("company", "")
        # Enlever les parenthèses vides ou avec seulement des espaces
        company_name = re.sub(r"\s*\(\s*\)\s*$", "", company_name).strip()
        # Enlever "(NON RENSEIGNÉ)" à la fin
        company_name = re.sub(
            r"\s*\(NON RENSEIGNÉ\)\s*$", "", company_name, flags=re.IGNORECASE
        ).strip()

        company_run = header_para.add_run(company_name.upper())
        company_run.font.size = Pt(14)
        company_run.font.color.rgb = self.COLOR_GOLD
        company_run.font.bold = True
        company_run.font.name = "Calibri"

        # Tabulation pour passer à droite
        header_para.add_run("\t")

        # Période (droite, doré, gras)
        period_run = header_para.add_run(experience.get("period", ""))
        period_run.font.size = Pt(14)
        period_run.font.color.rgb = self.COLOR_GOLD
        period_run.font.bold = True
        period_run.font.name = "Calibri"

        # Ajouter une ligne horizontale directement sous ce paragraphe
        pPr = header_para._element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")

        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "000000")

        pBdr.append(bottom)
        pPr.append(pBdr)

        # Titre du poste (14pt, italique, gras)
        title_para = self.doc.add_paragraph()
        title_para.paragraph_format.space_after = Pt(8)
        title_para.paragraph_format.keep_with_next = True
        title_run = title_para.add_run(experience.get("title", "").upper())
        title_run.font.size = Pt(14)
        title_run.font.color.rgb = self.COLOR_GRAY
        title_run.font.bold = True
        title_run.font.italic = True
        title_run.font.name = "Calibri"

        # Contexte
        if experience.get("context"):
            context_para = self.doc.add_paragraph()
            context_para.paragraph_format.keep_with_next = True
            context_label = context_para.add_run(self.labels["contexte"] + " : ")
            context_label.font.bold = True
            context_label.font.underline = True
            context_label.font.size = Pt(12)
            context_label.font.color.rgb = self.COLOR_DARK_BLUE
            context_label.font.name = "Calibri"

            context_text = context_para.add_run(experience["context"])
            context_text.font.size = Pt(12)
            context_text.font.name = "Calibri"

        # Activités
        if experience.get("activities"):
            activities_para = self.doc.add_paragraph()
            activities_para.paragraph_format.space_before = Pt(8)
            activities_para.paragraph_format.keep_with_next = True

            activities_label = activities_para.add_run(self.labels["activites"] + " :")
            activities_label.font.bold = True
            activities_label.font.underline = True
            activities_label.font.size = Pt(12)
            activities_label.font.color.rgb = self.COLOR_DARK_BLUE
            activities_label.font.name = "Calibri"

            self._add_bullet_list(experience["activities"])

        # Environnement technique
        if experience.get("tech_env"):
            tech_para = self.doc.add_paragraph()
            tech_para.paragraph_format.space_before = Pt(8)
            tech_para.paragraph_format.space_after = Pt(25)

            tech_label = tech_para.add_run(self.labels["env_tech"] + " : ")
            tech_label.font.bold = True
            tech_label.font.underline = True
            tech_label.font.size = Pt(12)
            tech_label.font.color.rgb = self.COLOR_DARK_BLUE
            tech_label.font.name = "Calibri"

            tech_text = tech_para.add_run(experience["tech_env"])
            tech_text.font.size = Pt(12)
            tech_text.font.italic = True
            tech_text.font.color.rgb = self.COLOR_GRAY
            tech_text.font.name = "Calibri"

    def _add_horizontal_line(self):
        """Ajoute une ligne horizontale"""
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(5)

        pPr = para._element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")

        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "000000")

        pBdr.append(bottom)
        pPr.append(pBdr)

    def _add_experiences(self, cv_data):
        """Ajoute la section Expériences professionnelles"""
        experiences = cv_data.get("experiences", [])

        if not experiences:
            return

        self._add_section_title(self.labels["experiences"])

        for experience in experiences:
            self._add_experience_block(experience)

    def generate(self, cv_data, output_path):
        """
        Génère le document DOCX complet

        Args:
            cv_data: Dictionnaire contenant les données du CV
            output_path: Chemin du fichier de sortie

        Returns:
            str: Chemin du fichier généré
        """
        # Ajouter l'en-tête de page
        self._add_page_header(cv_data, output_path)

        # Ajouter le pied de page
        self._add_page_footer()

        # Génération des sections
        self._add_header(cv_data)
        self._add_competences(cv_data)
        self._add_formations(cv_data)
        self._add_experiences(cv_data)

        # Sauvegarde du document
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        self.doc.save(output_path)
        print(f"✓ Fichier DOCX généré : {output_path}")

        return str(output_path)


def generate_docx_from_cv_data(cv_data, output_path, target_language="fr"):
    """
    Fonction utilitaire pour générer un DOCX à partir de données CV

    Args:
        cv_data: Dictionnaire contenant les données du CV
        output_path: Chemin du fichier de sortie
        target_language: Langue cible pour les labels (fr, en, it, es)

    Returns:
        str: Chemin du fichier généré
    """
    generator = CVDocxGenerator(target_language=target_language)
    return generator.generate(cv_data, output_path)


if __name__ == "__main__":
    # Test avec des données exemple
    test_data = {
        "header": {
            "name": "Jean Dupont",
            "title": "Développeur Full Stack",
            "experience": "10 ans d'expérience",
        },
        "competences": {
            "operationnelles": [
                "Développement d'applications web",
                "Architecture logicielle",
                "Gestion de projet agile",
            ],
            "techniques": [
                {"category": "Langages", "items": "Python, JavaScript, Java"},
                {"category": "Frameworks", "items": "Django, React, Spring Boot"},
            ],
        },
        "formations": [
            {"year": "2015", "description": "Master Informatique, Université de Paris"},
            {"year": "2013", "description": "Licence Informatique, Université de Lyon"},
        ],
        "experiences": [
            {
                "company": "Tech Corp (Paris)",
                "period": "2020 à aujourd'hui",
                "title": "Lead Developer",
                "context": "Développement d'une plateforme SaaS pour la gestion de projets",
                "activities": [
                    "Architecture et développement backend",
                    "Mentorat des développeurs juniors",
                    "Mise en place de l'intégration continue",
                ],
                "tech_env": "Python, Django, React, PostgreSQL, Docker, AWS",
            }
        ],
    }

    output_file = "test_cv.docx"
    print(f"Génération d'un CV de test : {output_file}")
    generate_docx_from_cv_data(test_data, output_file)
    print("Test terminé !")
